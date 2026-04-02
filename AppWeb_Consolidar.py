#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HACK14 DATA SUITE
=================
Aplicación profesional de análisis y consolidación de datos
con IA integrada para documentos Word, PDF y Excel/CSV.
"""

import os, re, io, time, traceback, warnings, tempfile
import streamlit as st
import pandas as pd

# ─────────────────────────────────────────────
#  PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="HACK14 · Data Suite",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
#  ESTILOS GLOBALES
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Mono:wght@400;700&family=Syne:wght@400;600;700;800&display=swap');

:root {
    --bg:       #070a10;
    --bg2:      #0d1117;
    --bg3:      #161b22;
    --border:   #21262d;
    --accent:   #00d9ff;
    --accent2:  #7c3aed;
    --gold:     #f59e0b;
    --green:    #10b981;
    --red:      #ef4444;
    --text:     #e6edf3;
    --muted:    #8b949e;
    --radius:   10px;
}

/* Reset global */
html, body, [class*="css"] {
    font-family: 'Syne', sans-serif !important;
    color: var(--text) !important;
}

/* Streamlit background */
.stApp { background: var(--bg) !important; }
.stSidebar > div:first-child { background: var(--bg2) !important; border-right: 1px solid var(--border); }
[data-testid="stSidebarContent"] { background: var(--bg2) !important; }

/* Header principal */
.hack14-header {
    background: linear-gradient(135deg, #0d1117 0%, #0f1923 50%, #070a10 100%);
    border: 1px solid var(--border);
    border-top: 3px solid var(--accent);
    border-radius: var(--radius);
    padding: 28px 36px;
    margin-bottom: 24px;
    position: relative;
    overflow: hidden;
}
.hack14-header::before {
    content: '';
    position: absolute;
    top: -50%;
    right: -10%;
    width: 400px;
    height: 400px;
    background: radial-gradient(circle, rgba(0,217,255,0.06) 0%, transparent 70%);
    pointer-events: none;
}
.hack14-header h1 {
    font-family: 'Space Mono', monospace !important;
    font-size: 2rem !important;
    font-weight: 700 !important;
    letter-spacing: -0.5px;
    color: var(--text) !important;
    margin: 0 0 4px !important;
}
.hack14-header h1 span { color: var(--accent); }
.hack14-header p {
    color: var(--muted) !important;
    font-size: 0.875rem !important;
    margin: 0 !important;
    font-family: 'Space Mono', monospace !important;
}
.badge {
    display: inline-block;
    background: rgba(0,217,255,0.12);
    color: var(--accent);
    border: 1px solid rgba(0,217,255,0.3);
    padding: 2px 10px;
    border-radius: 20px;
    font-size: 0.7rem;
    font-family: 'Space Mono', monospace;
    font-weight: 700;
    letter-spacing: 1px;
    text-transform: uppercase;
    margin-left: 10px;
    vertical-align: middle;
}

/* Cards */
.card {
    background: var(--bg3);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 20px 24px;
    margin-bottom: 16px;
    transition: border-color 0.2s;
}
.card:hover { border-color: rgba(0,217,255,0.3); }
.card-title {
    font-size: 0.75rem;
    font-family: 'Space Mono', monospace;
    color: var(--muted);
    text-transform: uppercase;
    letter-spacing: 2px;
    margin-bottom: 8px;
}
.card-value {
    font-size: 2rem;
    font-weight: 800;
    color: var(--accent);
    line-height: 1;
}
.card-sub { font-size: 0.8rem; color: var(--muted); margin-top: 4px; }

/* Doc card */
.doc-card {
    background: var(--bg3);
    border: 1px solid var(--border);
    border-left: 3px solid var(--accent2);
    border-radius: var(--radius);
    padding: 14px 18px;
    margin-bottom: 10px;
}
.doc-card.pdf { border-left-color: var(--red); }
.doc-card.docx { border-left-color: var(--accent); }
.doc-card.txt { border-left-color: var(--gold); }
.doc-title { font-weight: 700; font-size: 0.9rem; color: var(--text); }
.doc-meta { font-size: 0.75rem; color: var(--muted); font-family: 'Space Mono', monospace; }

/* Mensajes de chat */
.chat-msg-user {
    background: rgba(124,58,237,0.15);
    border: 1px solid rgba(124,58,237,0.3);
    border-radius: var(--radius) var(--radius) 0 var(--radius);
    padding: 12px 16px;
    margin: 8px 0 8px auto;
    max-width: 80%;
    font-size: 0.9rem;
}
.chat-msg-ai {
    background: rgba(0,217,255,0.08);
    border: 1px solid rgba(0,217,255,0.2);
    border-radius: var(--radius) var(--radius) var(--radius) 0;
    padding: 12px 16px;
    margin: 8px auto 8px 0;
    max-width: 90%;
    font-size: 0.9rem;
    line-height: 1.6;
}
.chat-label {
    font-size: 0.65rem;
    font-family: 'Space Mono', monospace;
    text-transform: uppercase;
    letter-spacing: 1px;
    color: var(--muted);
    margin-bottom: 4px;
}

/* Sidebar nav */
.nav-item {
    display: flex;
    align-items: center;
    gap: 10px;
    padding: 10px 14px;
    border-radius: 8px;
    cursor: pointer;
    transition: all 0.2s;
    margin-bottom: 4px;
    font-size: 0.875rem;
    font-weight: 600;
    color: var(--muted);
}
.nav-item:hover { background: var(--bg3); color: var(--text); }
.nav-item.active { background: rgba(0,217,255,0.1); color: var(--accent); border: 1px solid rgba(0,217,255,0.2); }

/* Botones Streamlit */
.stButton > button {
    background: var(--accent) !important;
    color: #000 !important;
    border: none !important;
    font-family: 'Space Mono', monospace !important;
    font-weight: 700 !important;
    font-size: 0.8rem !important;
    letter-spacing: 0.5px !important;
    border-radius: 6px !important;
    padding: 8px 20px !important;
    transition: all 0.2s !important;
}
.stButton > button:hover {
    background: #00b8d9 !important;
    transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(0,217,255,0.3) !important;
}

/* Inputs */
.stTextInput > div > div > input,
.stTextArea textarea,
.stSelectbox > div > div {
    background: var(--bg2) !important;
    border: 1px solid var(--border) !important;
    color: var(--text) !important;
    border-radius: 6px !important;
    font-family: 'Syne', sans-serif !important;
}
.stTextInput > div > div > input:focus,
.stTextArea textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 2px rgba(0,217,255,0.2) !important;
}

/* File uploader */
[data-testid="stFileUploadDropzone"] {
    background: var(--bg3) !important;
    border: 2px dashed var(--border) !important;
    border-radius: var(--radius) !important;
    transition: all 0.2s !important;
}
[data-testid="stFileUploadDropzone"]:hover {
    border-color: var(--accent) !important;
    background: rgba(0,217,255,0.04) !important;
}

/* Progress */
.stProgress > div > div > div { background: var(--accent) !important; }

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
    background: var(--bg2) !important;
    border-radius: 8px 8px 0 0 !important;
    border: 1px solid var(--border) !important;
    border-bottom: none !important;
    gap: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    font-family: 'Space Mono', monospace !important;
    font-size: 0.75rem !important;
    text-transform: uppercase !important;
    letter-spacing: 1px !important;
    color: var(--muted) !important;
    padding: 10px 20px !important;
}
.stTabs [aria-selected="true"] {
    background: rgba(0,217,255,0.1) !important;
    color: var(--accent) !important;
    border-bottom: 2px solid var(--accent) !important;
}

/* Status pills */
.pill-ok { background:rgba(16,185,129,.15); color:#10b981; border:1px solid rgba(16,185,129,.3); padding:2px 10px; border-radius:20px; font-size:.7rem; font-family:'Space Mono',monospace; }
.pill-err { background:rgba(239,68,68,.15); color:#ef4444; border:1px solid rgba(239,68,68,.3); padding:2px 10px; border-radius:20px; font-size:.7rem; font-family:'Space Mono',monospace; }
.pill-warn { background:rgba(245,158,11,.15); color:#f59e0b; border:1px solid rgba(245,158,11,.3); padding:2px 10px; border-radius:20px; font-size:.7rem; font-family:'Space Mono',monospace; }

/* Dataframe */
.stDataFrame { border-radius: var(--radius) !important; overflow: hidden !important; }

/* Divider */
.h14-divider {
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--border) 20%, var(--border) 80%, transparent);
    margin: 24px 0;
}

/* Sidebar logo */
.sidebar-logo {
    font-family: 'Space Mono', monospace;
    font-size: 1.3rem;
    font-weight: 700;
    color: var(--accent);
    letter-spacing: -1px;
    padding: 16px 8px 8px;
    border-bottom: 1px solid var(--border);
    margin-bottom: 16px;
}
.sidebar-logo span { color: var(--muted); font-size: 0.65rem; display: block; font-weight: 400; margin-top: 2px; letter-spacing: 2px; text-transform: uppercase; }

/* Summary box */
.summary-box {
    background: linear-gradient(135deg, rgba(0,217,255,0.06), rgba(124,58,237,0.06));
    border: 1px solid rgba(0,217,255,0.2);
    border-radius: var(--radius);
    padding: 18px 22px;
    font-size: 0.875rem;
    line-height: 1.7;
    color: var(--text);
}

/* Error box */
.error-box {
    background: rgba(239,68,68,0.08);
    border: 1px solid rgba(239,68,68,0.3);
    border-radius: var(--radius);
    padding: 12px 16px;
    color: var(--red);
    font-size: 0.85rem;
}

/* Scrollbar */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: var(--border); border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: var(--muted); }

/* Metric override */
[data-testid="metric-container"] {
    background: var(--bg3) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius) !important;
    padding: 16px !important;
}
[data-testid="stMetricValue"] { color: var(--accent) !important; font-family: 'Space Mono', monospace !important; }
[data-testid="stMetricLabel"] { color: var(--muted) !important; font-size: 0.75rem !important; }

/* Checkbox */
.stCheckbox label { color: var(--text) !important; }

/* Success / warning / error messages */
.stSuccess { background: rgba(16,185,129,0.1) !important; border: 1px solid rgba(16,185,129,0.3) !important; }
.stWarning { background: rgba(245,158,11,0.1) !important; border: 1px solid rgba(245,158,11,0.3) !important; }
.stError   { background: rgba(239,68,68,0.1)  !important; border: 1px solid rgba(239,68,68,0.3)  !important; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  IMPORTS OPCIONALES
# ─────────────────────────────────────────────
try:
    import anthropic
    _ANTHROPIC_OK = True
except ImportError:
    _ANTHROPIC_OK = False

try:
    import google.genai as genai_new
    import google.genai.types as genai_types
    _GEMINI_OK = True
    _GEMINI_SDK = "new"   # google-genai >= 1.0
except ImportError:
    try:
        import google.generativeai as genai_old
        _GEMINI_OK = True
        _GEMINI_SDK = "old"   # google-generativeai (legacy)
    except ImportError:
        _GEMINI_OK = False
        _GEMINI_SDK = None

try:
    from openai import OpenAI as OpenAIClient
    _OPENAI_OK = True
except ImportError:
    _OPENAI_OK = False

try:
    from docx import Document as DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from pypdf import PdfReader
    _PDF_OK = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        _PDF_OK = True
    except ImportError:
        _PDF_OK = False

try:
    import chardet
    _CHARDET_OK = True
except ImportError:
    _CHARDET_OK = False


# ─────────────────────────────────────────────
#  CATÁLOGO DE PROVEEDORES Y MODELOS (agrupados)
# ─────────────────────────────────────────────

# Cada modelo puede tener: id, label, ctx, best, group, icon_badge
AI_PROVIDERS = {
    "Google Gemini": {
        "icon": "🟦",
        "color": "#4285F4",
        "sdk": "gemini",
        "sdk_ok": lambda: _GEMINI_OK,
        "key_placeholder": "AIzaSy...",
        "console_url": "https://aistudio.google.com/app/apikey",
        "console_label": "aistudio.google.com",
        "model_groups": [
            {
                "group": "Gemini 2.5 (Última Generación)",
                "models": [
                    {"id": "gemini-2.5-pro",             "label": "gemini-2.5-pro",          "ctx": "1M",  "best": True,  "badge": "⭐"},
                    {"id": "gemini-2.5-flash",            "label": "gemini-2.5-flash",         "ctx": "1M",  "best": False, "badge": "⚡"},
                ]
            },
            {
                "group": "Gemini 2.0",
                "models": [
                    {"id": "gemini-2.0-pro-exp-02-05",   "label": "gemini-2.0-pro-exp-02-05", "ctx": "2M",  "best": False, "badge": "🔬"},
                    {"id": "gemini-2.0-flash",            "label": "gemini-2.0-flash",         "ctx": "1M",  "best": False, "badge": "⚡"},
                    {"id": "gemini-2.0-flash-lite",       "label": "gemini-2.0-flash-lite",    "ctx": "1M",  "best": False, "badge": "🪶"},
                ]
            },
            {
                "group": "Gemini 1.5 (Estables Anteriores)",
                "models": [
                    {"id": "gemini-1.5-pro",              "label": "gemini-1.5-pro",           "ctx": "2M",  "best": False, "badge": ""},
                    {"id": "gemini-1.5-flash",            "label": "gemini-1.5-flash",         "ctx": "1M",  "best": False, "badge": ""},
                    {"id": "gemini-1.5-flash-8b",         "label": "gemini-1.5-flash-8b",      "ctx": "1M",  "best": False, "badge": "🪶"},
                ]
            },
            {
                "group": "Personalizado",
                "models": [
                    {"id": "__custom__", "label": "Otro (Ingresar nombre manualmente...)", "ctx": "", "best": False, "badge": ""},
                ]
            },
        ]
    },
    "Anthropic Claude": {
        "icon": "🟠",
        "color": "#D97706",
        "sdk": "anthropic",
        "sdk_ok": lambda: _ANTHROPIC_OK,
        "key_placeholder": "sk-ant-api03-...",
        "console_url": "https://console.anthropic.com",
        "console_label": "console.anthropic.com",
        "model_groups": [
            {
                "group": "Claude 4 (Última Generación)",
                "models": [
                    {"id": "claude-opus-4-5",        "label": "Claude Opus 4.5",   "ctx": "200K", "best": True,  "badge": "⭐"},
                    {"id": "claude-sonnet-4-5",      "label": "Claude Sonnet 4.5", "ctx": "200K", "best": False, "badge": "⚡"},
                ]
            },
            {
                "group": "Claude 3.5",
                "models": [
                    {"id": "claude-3-5-sonnet-20241022", "label": "Claude 3.5 Sonnet", "ctx": "200K", "best": False, "badge": ""},
                    {"id": "claude-haiku-3-5",           "label": "Claude 3.5 Haiku",  "ctx": "200K", "best": False, "badge": "⚡"},
                ]
            },
            {
                "group": "Claude 3 (Estables)",
                "models": [
                    {"id": "claude-3-opus-20240229",  "label": "Claude 3 Opus",   "ctx": "200K", "best": False, "badge": ""},
                    {"id": "claude-3-haiku-20240307", "label": "Claude 3 Haiku",  "ctx": "200K", "best": False, "badge": ""},
                ]
            },
            {
                "group": "Personalizado",
                "models": [
                    {"id": "__custom__", "label": "Otro (Ingresar nombre manualmente...)", "ctx": "", "best": False, "badge": ""},
                ]
            },
        ]
    },
    "OpenAI / ChatGPT": {
        "icon": "🟢",
        "color": "#10A37F",
        "sdk": "openai",
        "sdk_ok": lambda: _OPENAI_OK,
        "key_placeholder": "sk-proj-...",
        "console_url": "https://platform.openai.com/api-keys",
        "console_label": "platform.openai.com",
        "model_groups": [
            {
                "group": "GPT-4o",
                "models": [
                    {"id": "gpt-4o",       "label": "gpt-4o",       "ctx": "128K", "best": True,  "badge": "⭐"},
                    {"id": "gpt-4o-mini",  "label": "gpt-4o-mini",  "ctx": "128K", "best": False, "badge": "⚡"},
                ]
            },
            {
                "group": "GPT-4",
                "models": [
                    {"id": "gpt-4-turbo",  "label": "gpt-4-turbo",  "ctx": "128K", "best": False, "badge": ""},
                    {"id": "gpt-4",        "label": "gpt-4",         "ctx": "8K",   "best": False, "badge": ""},
                ]
            },
            {
                "group": "o1 Razonamiento",
                "models": [
                    {"id": "o1-preview",   "label": "o1-preview",   "ctx": "128K", "best": False, "badge": "🔬"},
                    {"id": "o1-mini",      "label": "o1-mini",       "ctx": "128K", "best": False, "badge": "⚡"},
                ]
            },
            {
                "group": "GPT-3.5",
                "models": [
                    {"id": "gpt-3.5-turbo","label": "gpt-3.5-turbo","ctx": "16K",  "best": False, "badge": ""},
                ]
            },
            {
                "group": "Personalizado",
                "models": [
                    {"id": "__custom__", "label": "Otro (Ingresar nombre manualmente...)", "ctx": "", "best": False, "badge": ""},
                ]
            },
        ]
    },
}

def get_all_models(provider_name):
    """Retorna lista plana de todos los modelos de un proveedor."""
    groups = AI_PROVIDERS.get(provider_name, {}).get("model_groups", [])
    all_models = []
    for g in groups:
        all_models.extend(g["models"])
    return all_models

def get_best_model(provider_name):
    """Retorna el ID del modelo recomendado."""
    for m in get_all_models(provider_name):
        if m.get("best"):
            return m["id"]
    models = get_all_models(provider_name)
    return models[0]["id"] if models else ""


# ─────────────────────────────────────────────
#  ESTADO DE SESIÓN
# ─────────────────────────────────────────────
def init_state():
    defaults = {
        "page": "consolidador",
        "api_key": "",
        "ai_provider": "Google Gemini",
        "ai_model": "gemini-2.5-pro",
        "ai_model_custom": "",       # para el campo "Personalizado"
        "chat_history": [],
        "documentos_cargados": [],
        "doc_textos": {},
        "doc_resumenes": {},
        "consolidado_df": None,
        "errores_consol": [],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()


# ─────────────────────────────────────────────
#  LÓGICA DE LECTURA Y CONSOLIDACIÓN
# ─────────────────────────────────────────────
def normalizar_columnas(cols):
    return [str(c).strip() for c in cols]

def detectar_encoding(data_bytes):
    if _CHARDET_OK:
        r = chardet.detect(data_bytes[:20000])
        enc = r.get('encoding') or 'utf-8'
        if r.get('confidence', 1) < 0.6:
            enc = 'utf-8-sig'
        return enc
    return 'utf-8-sig'

def leer_csv_bytes(data_bytes, filename):
    enc = detectar_encoding(data_bytes)
    texto = data_bytes.decode(enc, errors='replace')
    from io import StringIO
    counts = {sep: texto[:5000].count(sep) for sep in [',', ';', '\t', '|']}
    sep = max(counts, key=counts.get)
    errores = []
    for s in [sep, ',', ';', '\t']:
        try:
            df = pd.read_csv(StringIO(texto), sep=s, dtype=str, engine='python')
            if not df.empty and len(df.columns) >= 2:
                df.columns = normalizar_columnas(df.columns.tolist())
                return [("CSV", df)]
        except Exception as e:
            errores.append(str(e))
    raise ValueError(f"No se pudo leer CSV '{filename}': {errores}")

def leer_excel_bytes(data_bytes, filename):
    buf = io.BytesIO(data_bytes)
    xls = pd.ExcelFile(buf)
    hojas = []
    for nombre in xls.sheet_names:
        try:
            df = pd.read_excel(xls, sheet_name=nombre, dtype=str)
            df.replace({'nan': None, 'NaN': None}, inplace=True)
            df.columns = normalizar_columnas(df.columns.tolist())
            if df.dropna(how='all').empty or len(df.columns) < 2:
                continue
            hojas.append((nombre, df))
        except Exception as e:
            hojas.append((nombre, e))
    return hojas

def leer_archivo(uploaded_file):
    data_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    name = uploaded_file.name.lower()
    if name.endswith(('.xlsx', '.xls', '.xlsm')):
        return leer_excel_bytes(data_bytes, uploaded_file.name)
    elif name.endswith('.csv'):
        return leer_csv_bytes(data_bytes, uploaded_file.name)
    raise ValueError(f"Extensión no soportada: {uploaded_file.name}")

def columnas_coinciden(cols_ref, cols_test):
    sr, st_ = set(cols_ref), set(cols_test)
    return {'coincide': sr == st_, 'faltantes': sorted(sr - st_), 'sobrantes': sorted(st_ - sr)}

def consolidar_archivos(archivos_upload, eliminar_dup, incl_traza):
    datos, errores = [], []
    estructura_ref = None
    progress_bar = st.progress(0)
    status_text  = st.empty()

    for i, f in enumerate(archivos_upload):
        nombre = f.name
        status_text.markdown(f"<span style='color:#8b949e;font-family:Space Mono,monospace;font-size:.8rem;'>▶ [{i+1}/{len(archivos_upload)}] {nombre}</span>", unsafe_allow_html=True)
        try:
            hojas = leer_archivo(f)
            hojas_incluidas = 0

            for nombre_hoja, df in hojas:
                if isinstance(df, Exception):
                    errores.append({'archivo': nombre, 'hoja': nombre_hoja, 'tipo': 'Error lectura', 'detalle': str(df)})
                    continue

                if estructura_ref is None:
                    estructura_ref = df.columns.tolist()

                comp = columnas_coinciden(estructura_ref, df.columns.tolist())
                if not comp['coincide']:
                    det = []
                    if comp['faltantes']: det.append(f"Faltan: {comp['faltantes']}")
                    if comp['sobrantes']: det.append(f"Sobran: {comp['sobrantes']}")
                    errores.append({'archivo': nombre, 'hoja': nombre_hoja, 'tipo': 'Estructura diferente', 'detalle': '; '.join(det)})
                    continue

                df = df[estructura_ref].copy()
                for col in df.columns:
                    if df[col].dtype == object:
                        df[col] = df[col].str.strip().replace({'': None, 'nan': None, 'NaN': None, 'NULL': None, 'null': None, 'N/A': None, '#N/A': None, '-': None})
                mask = df.isna().all(axis=1)
                df = df[~mask]

                if incl_traza:
                    df.insert(0, '__archivo__', nombre)
                    df.insert(1, '__hoja__', nombre_hoja)

                datos.append(df)
                hojas_incluidas += 1

            if hojas_incluidas == 0:
                errores.append({'archivo': nombre, 'hoja': '*', 'tipo': 'Archivo rechazado', 'detalle': 'Ninguna hoja coincide con la estructura'})
        except Exception as e:
            errores.append({'archivo': nombre, 'hoja': '*', 'tipo': 'Error general', 'detalle': str(e)})

        progress_bar.progress((i + 1) / len(archivos_upload))

    status_text.empty()
    progress_bar.empty()

    if not datos:
        return None, errores

    df_final = pd.concat(datos, ignore_index=True)

    if eliminar_dup:
        cols_subset = [c for c in (estructura_ref or []) if c in df_final.columns]
        if cols_subset:
            antes = len(df_final)
            df_final.drop_duplicates(subset=cols_subset, keep='first', inplace=True)
            n_dup = antes - len(df_final)
            if n_dup > 0:
                errores.append({'archivo': 'GLOBAL', 'hoja': '*', 'tipo': 'Duplicados eliminados', 'detalle': str(n_dup)})

    return df_final, errores


# ─────────────────────────────────────────────
#  LÓGICA DE DOCUMENTOS (PDF / DOCX / TXT)
# ─────────────────────────────────────────────
def extraer_texto_pdf(data_bytes):
    """Extrae texto de PDF con múltiples estrategias de fallback."""
    errores = []

    # Estrategia 1: pypdf
    if _PDF_OK:
        try:
            reader = PdfReader(io.BytesIO(data_bytes))
            partes = []
            for i, page in enumerate(reader.pages):
                try:
                    t = page.extract_text()
                    if t and t.strip():
                        partes.append(t)
                except Exception as ep:
                    errores.append(f"Página {i}: {ep}")
            texto = "\n".join(partes).strip()
            if texto:
                return texto
            errores.append("pypdf extrajo 0 caracteres (posible PDF escaneado o protegido)")
        except Exception as e:
            errores.append(f"pypdf falló: {e}")

    # Estrategia 2: pdfminer (si está instalado)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        texto = pdfminer_extract(io.BytesIO(data_bytes))
        if texto and texto.strip():
            return texto.strip()
        errores.append("pdfminer extrajo 0 caracteres")
    except ImportError:
        pass
    except Exception as e:
        errores.append(f"pdfminer falló: {e}")

    # Estrategia 3: pypdfium2
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(io.BytesIO(data_bytes))
        partes = []
        for page in pdf:
            textpage = page.get_textpage()
            t = textpage.get_text_range()
            if t and t.strip():
                partes.append(t)
        texto = "\n".join(partes).strip()
        if texto:
            return texto
        errores.append("pypdfium2 extrajo 0 caracteres")
    except ImportError:
        pass
    except Exception as e:
        errores.append(f"pypdfium2 falló: {e}")

    # Si llegamos aquí, el PDF no tiene texto extraíble (es imagen)
    msg = (
        "Este PDF no contiene texto seleccionable (posiblemente es un escaneado o imagen). "
        f"Detalles técnicos: {' | '.join(errores) if errores else 'sin texto detectado'}"
    )
    raise ValueError(msg)

def extraer_texto_docx(data_bytes):
    if not _DOCX_OK:
        raise ImportError("python-docx no disponible")
    doc = DocxDocument(io.BytesIO(data_bytes))
    parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(parrafos)

def extraer_texto_txt(data_bytes):
    enc = detectar_encoding(data_bytes)
    return data_bytes.decode(enc, errors='replace')

def extraer_texto(nombre, data_bytes):
    ext = nombre.lower().rsplit('.', 1)[-1]
    if ext == 'pdf':
        return extraer_texto_pdf(data_bytes)
    elif ext in ('docx', 'doc'):
        return extraer_texto_docx(data_bytes)
    elif ext == 'txt':
        return extraer_texto_txt(data_bytes)
    raise ValueError(f"Extensión no soportada para análisis IA: .{ext}")

def tipo_doc(nombre):
    ext = nombre.lower().rsplit('.', 1)[-1]
    return {'pdf': 'pdf', 'docx': 'docx', 'doc': 'docx', 'txt': 'txt'}.get(ext, 'txt')


# ─────────────────────────────────────────────
#  CLIENTE IA — MULTI-PROVEEDOR
# ─────────────────────────────────────────────
def resolver_model_id():
    """Retorna el model ID real a usar (resuelve __custom__)."""
    mid = st.session_state.get("ai_model", "")
    if mid == "__custom__":
        return st.session_state.get("ai_model_custom", "").strip()
    return mid

def llamar_ia(system_prompt, user_prompt, max_tokens=1500):
    key      = st.session_state.get("api_key", "").strip()
    provider = st.session_state.get("ai_provider", "Google Gemini")
    model    = resolver_model_id()

    if not key:
        return None, "No hay API key configurada. Ve a ⚙️ Configuración."
    if not model:
        return None, "No hay modelo seleccionado. Ve a ⚙️ Configuración."

    sdk = AI_PROVIDERS.get(provider, {}).get("sdk", "")

    # ── GEMINI (nuevo SDK google-genai) ─────
    if sdk == "gemini":
        if not _GEMINI_OK:
            return None, "Instala el SDK de Gemini:\n  pip install google-genai"
        try:
            if _GEMINI_SDK == "new":
                client = genai_new.Client(api_key=key)
                resp = client.models.generate_content(
                    model=model,
                    contents=user_prompt,
                    config=genai_types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        max_output_tokens=max_tokens,
                    )
                )
                return resp.text, None
            else:
                # Legacy fallback
                import warnings
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore")
                    genai_old.configure(api_key=key)
                    m = genai_old.GenerativeModel(
                        model_name=model,
                        system_instruction=system_prompt
                    )
                    resp = m.generate_content(user_prompt)
                    return resp.text, None
        except Exception as e:
            err_str = str(e)
            # Mensajes de error más amigables
            if "API_KEY_INVALID" in err_str or "API key not valid" in err_str:
                return None, "❌ API Key inválida. Verifica que sea correcta en Google AI Studio."
            if "not found" in err_str.lower() or "404" in err_str:
                return None, f"❌ Modelo '{model}' no encontrado. Verifica el nombre exacto."
            if "quota" in err_str.lower():
                return None, "❌ Cuota de API agotada. Revisa tu plan en Google AI Studio."
            return None, f"❌ Error Gemini: {err_str}"

    # ── ANTHROPIC ───────────────────────────
    elif sdk == "anthropic":
        if not _ANTHROPIC_OK:
            return None, "Instala el SDK:\n  pip install anthropic"
        try:
            client = anthropic.Anthropic(api_key=key)
            msg = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}]
            )
            return msg.content[0].text, None
        except Exception as e:
            err_str = str(e)
            if "authentication" in err_str.lower() or "401" in err_str:
                return None, "❌ API Key inválida para Anthropic."
            if "not found" in err_str.lower() or "404" in err_str:
                return None, f"❌ Modelo '{model}' no encontrado."
            return None, f"❌ Error Anthropic: {err_str}"

    # ── OPENAI ──────────────────────────────
    elif sdk == "openai":
        if not _OPENAI_OK:
            return None, "Instala el SDK:\n  pip install openai"
        try:
            client = OpenAIClient(api_key=key)
            resp = client.chat.completions.create(
                model=model,
                max_tokens=max_tokens,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": user_prompt}
                ]
            )
            return resp.choices[0].message.content, None
        except Exception as e:
            err_str = str(e)
            if "401" in err_str or "Incorrect API key" in err_str:
                return None, "❌ API Key inválida para OpenAI."
            if "not found" in err_str.lower() or "404" in err_str:
                return None, f"❌ Modelo '{model}' no encontrado."
            return None, f"❌ Error OpenAI: {err_str}"

    return None, f"Proveedor desconocido: {provider}"

def generar_resumen(nombre, texto):
    system = (
        "Eres un analista de documentos experto. "
        "Genera resúmenes claros, estructurados y concisos. "
        "Usa formato Markdown. Sé preciso y útil."
    )
    prompt = (
        f"Analiza el siguiente documento llamado '{nombre}' y genera:\n"
        "1. **Resumen ejecutivo** (3-5 oraciones)\n"
        "2. **Temas principales** (lista de 3-7 puntos)\n"
        "3. **Datos clave o cifras** mencionados (si aplica)\n"
        "4. **Conclusiones o hallazgos** principales\n\n"
        f"---\nCONTENIDO DEL DOCUMENTO:\n{texto[:8000]}\n---"
    )
    return llamar_ia(system, prompt, max_tokens=1200)

def responder_pregunta(pregunta, documentos_ctx, historial):
    system = (
        "Eres un asistente experto en análisis de documentos. "
        "Respondes preguntas basándote ÚNICAMENTE en el contenido de los documentos proporcionados. "
        "Si la información no está en los documentos, lo dices claramente. "
        "Usa Markdown para dar respuestas claras y bien estructuradas."
    )
    ctx_doc = "\n\n".join(
        f"=== DOCUMENTO: {nombre} ===\n{texto[:4000]}"
        for nombre, texto in documentos_ctx.items()
    )
    hist_str = ""
    for h in historial[-6:]:
        rol = "Usuario" if h['role'] == 'user' else "Asistente"
        hist_str += f"{rol}: {h['content']}\n"

    prompt = (
        f"DOCUMENTOS DISPONIBLES:\n{ctx_doc}\n\n"
        f"HISTORIAL RECIENTE:\n{hist_str}\n"
        f"PREGUNTA ACTUAL: {pregunta}"
    )
    return llamar_ia(system, prompt, max_tokens=1500)


# ─────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown('<div class="sidebar-logo">⚡ HACK14<span>Data Suite v1.0</span></div>', unsafe_allow_html=True)

    st.markdown("**Navegación**")
    pages = [
        ("consolidador", "📊", "Consolidador Excel/CSV"),
        ("documentos",   "📄", "Análisis de Documentos"),
        ("config",       "⚙️", "Configuración IA"),
    ]
    for pid, icon, label in pages:
        active = "active" if st.session_state.page == pid else ""
        if st.button(f"{icon}  {label}", key=f"nav_{pid}", use_container_width=True):
            st.session_state.page = pid
            st.rerun()

    st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)

    # Estado IA
    api_ok = bool(st.session_state.api_key.strip())
    prov   = st.session_state.get("ai_provider", "Google Gemini")
    prov_info = AI_PROVIDERS.get(prov, {})
    prov_icon = prov_info.get("icon", "🤖")
    st.markdown(
        f'<div style="font-size:.75rem;font-family:Space Mono,monospace;color:var(--muted);">ESTADO IA &nbsp;'
        f'<span class="{"pill-ok" if api_ok else "pill-err"}">{"● ACTIVA" if api_ok else "● SIN KEY"}</span></div>',
        unsafe_allow_html=True
    )
    if api_ok:
        all_m = get_all_models(prov)
        model_label = next((m["label"] for m in all_m if m["id"] == st.session_state.get("ai_model")), st.session_state.get("ai_model",""))
        if st.session_state.get("ai_model") == "__custom__":
            model_label = st.session_state.get("ai_model_custom", "") or "Personalizado"
        st.markdown(f'<div style="font-size:.7rem;color:#8b949e;margin-top:4px;font-family:Space Mono,monospace;">{prov_icon} {prov} · {model_label}</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div style="font-size:.72rem;color:#8b949e;margin-top:6px;">Configura tu API key en ⚙️ para usar la IA</div>', unsafe_allow_html=True)

    n_docs = len(st.session_state.documentos_cargados)
    if n_docs:
        st.markdown(f'<div style="font-size:.75rem;font-family:Space Mono,monospace;color:var(--muted);margin-top:8px;">DOCS CARGADOS &nbsp;<span class="pill-ok">● {n_docs}</span></div>', unsafe_allow_html=True)

    if st.session_state.consolidado_df is not None:
        df_c = st.session_state.consolidado_df
        st.markdown(f'<div style="font-size:.75rem;font-family:Space Mono,monospace;color:var(--muted);margin-top:8px;">CONSOLIDADO &nbsp;<span class="pill-ok">● {len(df_c):,} filas</span></div>', unsafe_allow_html=True)

    st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div style="font-size:.65rem;color:#3d444d;font-family:Space Mono,monospace;text-align:center;">HACK14 © 2025<br>Todos los derechos reservados</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  HEADER GLOBAL
# ─────────────────────────────────────────────
page_titles = {
    "consolidador": ("📊 CONSOLIDADOR", "Unifica archivos Excel, CSV y datos tabulares"),
    "documentos":   ("📄 ANÁLISIS IA", "Lee, resume y consulta documentos Word y PDF"),
    "config":       ("⚙️ CONFIGURACIÓN", "Configura tu API key de Anthropic para activar la IA"),
}
titulo, subtitulo = page_titles.get(st.session_state.page, ("HACK14", ""))
st.markdown(f"""
<div class="hack14-header">
  <h1>⚡ HACK<span>14</span> <span class="badge">DATA SUITE</span></h1>
  <p>{titulo} · {subtitulo}</p>
</div>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
#  PÁGINA: CONFIGURACIÓN IA
# ═══════════════════════════════════════════════════════
if st.session_state.page == "config":

    # ── Selector de proveedor ──────────────────────────
    st.markdown('<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;text-transform:uppercase;letter-spacing:2px;margin-bottom:14px;">① Elige tu proveedor de IA</div>', unsafe_allow_html=True)

    provider_cols = st.columns(len(AI_PROVIDERS))
    for idx, (pname, pdata) in enumerate(AI_PROVIDERS.items()):
        with provider_cols[idx]:
            is_active = st.session_state.ai_provider == pname
            border_color = pdata["color"] if is_active else "#21262d"
            bg_color     = f"rgba({int(pdata['color'][1:3],16)},{int(pdata['color'][3:5],16)},{int(pdata['color'][5:7],16)},0.08)" if is_active else "var(--bg3)"
            sdk_name     = pdata["sdk"]
            sdk_available = pdata["sdk_ok"]()

            avail_html = '<span class="pill-ok">● Instalado</span>' if sdk_available else '<span class="pill-warn">● pip install requerido</span>'

            st.markdown(f"""
            <div style="background:{bg_color};border:2px solid {border_color};border-radius:10px;padding:18px 16px;text-align:center;cursor:pointer;transition:all .2s;">
              <div style="font-size:2rem;margin-bottom:6px;">{pdata['icon']}</div>
              <div style="font-weight:700;font-size:.95rem;color:{'#e6edf3' if is_active else '#8b949e'};">{pname}</div>
              <div style="font-size:.68rem;margin-top:6px;">{avail_html}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"{'✅ Seleccionado' if is_active else 'Seleccionar'}", key=f"sel_{pname}", use_container_width=True):
                st.session_state.ai_provider = pname
                st.session_state.ai_model = get_best_model(pname)
                st.session_state.ai_model_custom = ""
                st.session_state.api_key  = ""
                st.rerun()

    st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)

    # ── Config del proveedor seleccionado ─────────────
    prov      = st.session_state.ai_provider
    prov_data = AI_PROVIDERS[prov]
    sdk_ok    = prov_data["sdk_ok"]()

    col1, col2 = st.columns([2, 1])

    with col1:
        # SDK warning
        if not sdk_ok:
            sdk_pkg = {"gemini": "google-generativeai", "anthropic": "anthropic", "openai": "openai"}.get(prov_data["sdk"], "")
            st.markdown(f"""
            <div class="card" style="border-color:rgba(245,158,11,0.4);background:rgba(245,158,11,0.05);">
              <div style="font-weight:700;color:#f59e0b;margin-bottom:6px;">⚠️ SDK no instalado</div>
              <div style="font-size:.85rem;color:#8b949e;">Ejecuta en tu terminal:<br>
              <code style="background:#0d1117;padding:4px 8px;border-radius:4px;color:#00d9ff;">pip install {sdk_pkg}</code>
              </div>
            </div>
            """, unsafe_allow_html=True)

        # ── Selector de modelo agrupado ──────────────────
        st.markdown(f'<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;text-transform:uppercase;letter-spacing:2px;margin-bottom:10px;">② Modelo — {prov_data["icon"]} {prov}</div>', unsafe_allow_html=True)

        model_groups = prov_data.get("model_groups", [])

        # Construimos lista plana con separadores de grupo
        flat_options   = []   # lo que recibe st.selectbox (IDs únicos)
        flat_labels    = []   # etiqueta que muestra format_func
        flat_disabled  = []   # True si es encabezado de grupo

        for grp in model_groups:
            # Separador / encabezado de grupo (no seleccionable → usamos ID especial)
            sep_id = f"__sep__{grp['group']}"
            flat_options.append(sep_id)
            flat_labels.append(f"── {grp['group']}")
            flat_disabled.append(True)

            for m in grp["models"]:
                flat_options.append(m["id"])
                badge = f" {m['badge']}" if m.get("badge") else ""
                ctx   = f"  [{m['ctx']}]" if m.get("ctx") else ""
                flat_labels.append(f"  {m['label']}{badge}{ctx}")
                flat_disabled.append(False)

        # Índice actual (ignorar separadores)
        current_model = st.session_state.get("ai_model", flat_options[1] if len(flat_options) > 1 else "")
        try:
            current_idx = flat_options.index(current_model)
        except ValueError:
            current_idx = next((i for i, o in enumerate(flat_options) if not o.startswith("__sep__")), 1)

        # selectbox con format_func
        sel_idx = st.selectbox(
            "Modelo",
            options=range(len(flat_options)),
            format_func=lambda i: flat_labels[i],
            index=current_idx,
            label_visibility="collapsed",
            key="model_selectbox"
        )

        selected_id = flat_options[sel_idx]

        # Si se eligió un separador → volver al anterior válido
        if selected_id.startswith("__sep__"):
            selected_id = current_model

        st.session_state.ai_model = selected_id

        # Campo texto si es personalizado
        if selected_id == "__custom__":
            custom_val = st.text_input(
                "Nombre del modelo",
                value=st.session_state.get("ai_model_custom", ""),
                placeholder="Ej: gemini-2.5-pro-latest",
                label_visibility="visible"
            )
            st.session_state.ai_model_custom = custom_val.strip()
            display_model_id = custom_val.strip() or "(sin especificar)"
        else:
            st.session_state.ai_model_custom = ""
            display_model_id = selected_id

        # Tarjeta del modelo seleccionado
        # Buscar metadata del modelo
        all_flat = get_all_models(prov)
        m_meta   = next((m for m in all_flat if m["id"] == selected_id), None)
        m_label  = m_meta["label"] if m_meta else display_model_id
        m_ctx    = m_meta["ctx"] if m_meta else "—"
        m_badge  = m_meta.get("badge", "") if m_meta else ""

        st.markdown(f"""
        <div class="card" style="border-left:3px solid {prov_data['color']};margin-top:8px;">
          <div class="card-title">MODELO ACTIVO</div>
          <div style="font-weight:700;font-size:1rem;color:#e6edf3;">{prov_data['icon']} {m_label} {m_badge}</div>
          <div style="font-size:.75rem;color:#8b949e;font-family:Space Mono,monospace;margin-top:4px;">
            ID: <span style="color:#00d9ff;">{display_model_id}</span>
            {"&nbsp;&nbsp;|&nbsp;&nbsp;Contexto: " + m_ctx + " tokens" if m_ctx else ""}
          </div>
        </div>
        """, unsafe_allow_html=True)

        # API Key
        st.markdown('<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;text-transform:uppercase;letter-spacing:2px;margin:16px 0 8px;">③ API Key</div>', unsafe_allow_html=True)

        api_input = st.text_input(
            "API Key",
            value=st.session_state.api_key,
            type="password",
            placeholder=prov_data["key_placeholder"],
            label_visibility="collapsed"
        )
        if api_input != st.session_state.api_key:
            st.session_state.api_key = api_input

        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("💾 Guardar y verificar", use_container_width=True):
                real_model = resolver_model_id()
                if not api_input.strip():
                    st.warning("Ingresa una API key primero.")
                elif not sdk_ok:
                    sdk_pkg = {"gemini": "google-genai", "anthropic": "anthropic", "openai": "openai"}.get(prov_data["sdk"], "")
                    st.error(f"Instala el SDK primero: pip install {sdk_pkg}")
                elif not real_model:
                    st.warning("Escribe el nombre del modelo en el campo 'Personalizado'.")
                else:
                    with st.spinner(f"Verificando conexión con {prov} · {real_model}..."):
                        resp, err = llamar_ia("Responde solo con la palabra OK.", "test de conexión", max_tokens=20)
                    if err:
                        st.error(err)
                    else:
                        st.success(f"✅ Conexión verificada · {prov} · {real_model}")
        with col_b:
            if st.button("🗑️ Limpiar key", use_container_width=True):
                st.session_state.api_key = ""
                st.rerun()

    with col2:
        # Info del proveedor
        st.markdown(f"""
        <div class="card" style="border-left:3px solid {prov_data['color']};">
          <div class="card-title">🔑 Obtener API Key</div>
          <p style="font-size:.8rem;color:#8b949e;margin:0;">
            <strong style="color:#e6edf3;">{prov}</strong><br>
            Visita <a href="{prov_data['console_url']}" target="_blank" style="color:#00d9ff;">{prov_data['console_label']}</a>
            para obtener tu clave.<br><br>
            <strong style="color:#e6edf3;">Privacidad</strong><br>
            Tu key solo existe en esta sesión local. No se almacena en ningún servidor.
          </p>
        </div>
        """, unsafe_allow_html=True)

        # Panel de modelos agrupados (solo visual)
        panel_rows = ""
        for grp in prov_data.get("model_groups", []):
            panel_rows += f'<div style="font-size:.65rem;color:#4d5566;font-family:Space Mono,monospace;padding:6px 0 2px;text-transform:uppercase;letter-spacing:1px;">— {grp["group"]}</div>'
            for m in grp["models"]:
                if m["id"] == "__custom__":
                    continue
                is_sel = (m["id"] == st.session_state.get("ai_model"))
                badge  = m.get("badge", "")
                color  = "#e6edf3" if is_sel else "#8b949e"
                weight = "700" if is_sel else "400"
                prefix = "▶ " if is_sel else "  "
                panel_rows += (
                    f'<div style="display:flex;justify-content:space-between;align-items:center;'
                    f'padding:3px 4px;border-radius:4px;{"background:rgba(0,217,255,0.06);" if is_sel else ""}">'
                    f'<span style="font-size:.76rem;color:{color};font-weight:{weight};">{prefix}{m["label"]} {badge}</span>'
                    f'<span style="font-size:.65rem;color:#3d444d;font-family:Space Mono,monospace;">{m["ctx"]}</span>'
                    f'</div>'
                )

        st.markdown(f"""
        <div class="card" style="margin-top:12px;">
          <div class="card-title">📊 Modelos — {prov}</div>
          {panel_rows}
        </div>
        """, unsafe_allow_html=True)

        st.markdown("""
        <div class="card" style="border-left:3px solid var(--gold);margin-top:12px;">
          <div class="card-title">🔧 Funciones IA</div>
          <p style="font-size:.78rem;color:#8b949e;margin:0;">
            ✅ Resumen automático de documentos<br>
            ✅ Chat interactivo multi-documento<br>
            ✅ Preguntas en lenguaje natural<br>
            ✅ Extracción de datos clave<br>
            ✅ Soporte Gemini · Claude · GPT
          </p>
        </div>
        """, unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
#  PÁGINA: CONSOLIDADOR
# ═══════════════════════════════════════════════════════
elif st.session_state.page == "consolidador":

    tab1, tab2, tab3 = st.tabs(["📁 CARGAR ARCHIVOS", "📊 RESULTADO", "⚠️ ERRORES"])

    with tab1:
        col1, col2 = st.columns([3, 1])
        with col1:
            uploaded = st.file_uploader(
                "Arrastra aquí tus archivos Excel o CSV",
                type=['xlsx', 'xls', 'xlsm', 'csv'],
                accept_multiple_files=True,
                label_visibility="collapsed",
                help="Soporta .xlsx, .xls, .xlsm y .csv"
            )

        with col2:
            st.markdown('<div class="card"><div class="card-title">⚙️ Opciones</div></div>', unsafe_allow_html=True)
            eliminar_dup = st.checkbox("Eliminar duplicados", value=True)
            incl_traza   = st.checkbox("Columnas de trazabilidad", value=True)
            nombre_sal   = st.text_input("Nombre del archivo de salida", value="consolidado_hack14", label_visibility="visible")

        if uploaded:
            n = len(uploaded)
            cols = st.columns(4)
            cols[0].metric("Archivos cargados", n)
            cols[1].metric("Excel/XLSM", sum(1 for f in uploaded if f.name.lower().endswith(('.xlsx','.xls','.xlsm'))))
            cols[2].metric("CSV", sum(1 for f in uploaded if f.name.lower().endswith('.csv')))
            total_mb = sum(f.size for f in uploaded) / 1024 / 1024
            cols[3].metric("Tamaño total", f"{total_mb:.2f} MB")

            st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)

            # Vista previa archivos
            st.markdown('<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;text-transform:uppercase;letter-spacing:2px;margin-bottom:12px;">Archivos a procesar</div>', unsafe_allow_html=True)
            for f in uploaded:
                ext = f.name.lower().rsplit('.', 1)[-1]
                color = {'xlsx':'var(--green)','xls':'var(--green)','xlsm':'var(--green)','csv':'var(--gold)'}.get(ext, 'var(--accent)')
                size_kb = f.size / 1024
                size_str = f"{size_kb:.1f} KB" if size_kb < 1024 else f"{size_kb/1024:.2f} MB"
                st.markdown(f"""
                <div class="card" style="padding:12px 18px;border-left:3px solid {color};">
                  <span style="font-weight:700;font-size:.875rem;">{f.name}</span>
                  <span style="float:right;font-family:Space Mono,monospace;font-size:.72rem;color:#8b949e;">{size_str} &nbsp;|&nbsp; .{ext.upper()}</span>
                </div>
                """, unsafe_allow_html=True)

            st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)
            if st.button("🚀 INICIAR CONSOLIDACIÓN", use_container_width=True):
                with st.spinner("Consolidando archivos..."):
                    df_final, errores = consolidar_archivos(uploaded, eliminar_dup, incl_traza)
                st.session_state.consolidado_df = df_final
                st.session_state.errores_consol = errores
                if df_final is not None:
                    st.success(f"✅ Consolidación completada: {len(df_final):,} filas × {len(df_final.columns)} columnas")
                    st.session_state.page = "consolidador"
                    st.rerun()
                else:
                    st.error("❌ No se generaron datos. Revisa la pestaña de Errores.")
        else:
            st.markdown("""
            <div class="card" style="text-align:center;padding:48px;border-style:dashed;">
              <div style="font-size:3rem;margin-bottom:12px;">📂</div>
              <div style="font-size:1.1rem;font-weight:700;color:#e6edf3;">Ningún archivo cargado</div>
              <div style="font-size:.85rem;color:#8b949e;margin-top:8px;">
                Arrastra archivos .xlsx, .xls, .xlsm o .csv al área de carga superior
              </div>
            </div>
            """, unsafe_allow_html=True)

    with tab2:
        df = st.session_state.consolidado_df
        if df is not None and not df.empty:
            cols = st.columns(4)
            cols[0].metric("Total filas", f"{len(df):,}")
            cols[1].metric("Columnas", len(df.columns))
            cols[2].metric("Valores nulos", f"{df.isna().sum().sum():,}")
            dup_count = df.duplicated().sum()
            cols[3].metric("Duplicados", f"{dup_count:,}")

            st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)

            # Preview
            st.markdown('<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;text-transform:uppercase;letter-spacing:2px;margin-bottom:8px;">Vista previa (primeras 200 filas)</div>', unsafe_allow_html=True)
            st.dataframe(df.head(200), use_container_width=True, height=380)

            st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)

            # Descargas
            col_a, col_b = st.columns(2)
            with col_a:
                csv_bytes = df.to_csv(index=False, encoding='utf-8-sig').encode('utf-8-sig')
                st.download_button(
                    label="💾 Descargar CSV",
                    data=csv_bytes,
                    file_name=f"{nombre_sal if 'nombre_sal' in dir() else 'consolidado_hack14'}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            with col_b:
                buf_xlsx = io.BytesIO()
                try:
                    df.to_excel(buf_xlsx, index=False, engine='openpyxl')
                    buf_xlsx.seek(0)
                    st.download_button(
                        label="📊 Descargar Excel",
                        data=buf_xlsx.getvalue(),
                        file_name=f"{nombre_sal if 'nombre_sal' in dir() else 'consolidado_hack14'}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                except Exception as e:
                    st.warning(f"No se puede generar Excel: {e}")

            # Estadísticas columnas
            st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)
            st.markdown('<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;text-transform:uppercase;letter-spacing:2px;margin-bottom:8px;">Diagnóstico de columnas</div>', unsafe_allow_html=True)

            stats_rows = []
            for col in df.columns:
                nulos = int(df[col].isna().sum())
                pct   = nulos / max(len(df), 1) * 100
                uniq  = df[col].nunique()
                stats_rows.append({
                    "Columna": col,
                    "Tipo": str(df[col].dtype),
                    "Nulos": nulos,
                    "% Nulo": f"{pct:.1f}%",
                    "Únicos": uniq,
                    "Estado": "⚠️ Alerta" if pct > 50 else ("✅ OK" if pct < 10 else "🔶 Revisar")
                })
            st.dataframe(pd.DataFrame(stats_rows), use_container_width=True)
        else:
            st.markdown("""
            <div class="card" style="text-align:center;padding:48px;">
              <div style="font-size:2rem;margin-bottom:12px;">📊</div>
              <div style="color:#8b949e;font-size:.9rem;">Ejecuta la consolidación para ver los resultados aquí</div>
            </div>
            """, unsafe_allow_html=True)

    with tab3:
        errores = st.session_state.errores_consol
        if errores:
            tipos = list(set(e['tipo'] for e in errores))
            filtro = st.selectbox("Filtrar por tipo", ["Todos"] + sorted(tipos), label_visibility="visible")
            errores_f = errores if filtro == "Todos" else [e for e in errores if e['tipo'] == filtro]

            st.markdown(f'<div style="margin-bottom:12px;"><span class="pill-err">● {len(errores_f)} registros</span></div>', unsafe_allow_html=True)
            st.dataframe(pd.DataFrame(errores_f), use_container_width=True)

            err_csv = pd.DataFrame(errores).to_csv(index=False, encoding='utf-8-sig').encode('utf-8-sig')
            st.download_button("📥 Descargar reporte de errores", err_csv, file_name="errores_hack14.csv", mime="text/csv")
        else:
            st.markdown("""
            <div class="card" style="text-align:center;padding:32px;border-color:rgba(16,185,129,0.3);">
              <div style="font-size:2rem;margin-bottom:8px;">✅</div>
              <div style="color:#10b981;">Sin errores registrados</div>
            </div>
            """, unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
#  PÁGINA: ANÁLISIS DE DOCUMENTOS
# ═══════════════════════════════════════════════════════
elif st.session_state.page == "documentos":
    api_ok = bool(st.session_state.api_key.strip())

    if not api_ok:
        st.markdown("""
        <div class="card" style="border-color:rgba(245,158,11,0.4);background:rgba(245,158,11,0.05);">
          <div style="font-size:1rem;font-weight:700;color:#f59e0b;margin-bottom:6px;">⚠️ API Key no configurada</div>
          <div style="font-size:.85rem;color:#8b949e;">
            Para usar el análisis inteligente, configura tu API key de Anthropic en 
            <strong style="color:#e6edf3;">⚙️ Configuración</strong>.<br>
            Igual puedes cargar documentos y extraer su texto sin IA.
          </div>
        </div>
        """, unsafe_allow_html=True)

    col_izq, col_der = st.columns([2, 3])

    # ─── Panel izquierdo: carga y lista de docs ───
    with col_izq:
        st.markdown('<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;text-transform:uppercase;letter-spacing:2px;margin-bottom:10px;">Cargar documentos</div>', unsafe_allow_html=True)

        nuevos_docs = st.file_uploader(
            "Sube documentos",
            type=['pdf', 'docx', 'doc', 'txt'],
            accept_multiple_files=True,
            label_visibility="collapsed",
            help="Soporta PDF, Word (.docx) y texto plano (.txt)"
        )

        if nuevos_docs:
            nombres_ya_cargados = {d['nombre'] for d in st.session_state.documentos_cargados}
            hay_nuevos = False
            for f in nuevos_docs:
                if f.name in nombres_ya_cargados:
                    continue
                hay_nuevos = True
                # Leer bytes INMEDIATAMENTE antes de cualquier otra operación
                try:
                    f.seek(0)
                    data = f.read()
                    if not data:
                        st.error(f"❌ {f.name}: el archivo está vacío o no se pudo leer.")
                        continue
                except Exception as e:
                    st.error(f"❌ {f.name}: error al leer el archivo — {e}")
                    continue

                with st.spinner(f"Procesando {f.name}..."):
                    try:
                        texto = extraer_texto(f.name, data)
                        n_chars = len(texto)
                        st.session_state.documentos_cargados.append({
                            'nombre': f.name,
                            'tipo': tipo_doc(f.name),
                            'size': len(data),
                            'chars': n_chars,
                        })
                        st.session_state.doc_textos[f.name] = texto
                        if n_chars > 0:
                            st.success(f"✅ {f.name} — {n_chars:,} caracteres extraídos")
                        else:
                            st.warning(f"⚠️ {f.name} cargado pero sin texto extraíble.")
                    except Exception as e:
                        st.error(f"❌ {f.name}: {e}")
                        # Guardar igual para que el usuario vea el error en la lista
                        st.session_state.documentos_cargados.append({
                            'nombre': f.name,
                            'tipo': tipo_doc(f.name),
                            'size': len(data),
                            'chars': 0,
                            'error': str(e),
                        })
                        st.session_state.doc_textos[f.name] = ""
            if hay_nuevos:
                st.rerun()

        st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)

        # Lista de documentos cargados
        docs = st.session_state.documentos_cargados
        if docs:
            st.markdown(f'<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;text-transform:uppercase;letter-spacing:2px;margin-bottom:10px;">Documentos ({len(docs)})</div>', unsafe_allow_html=True)
            for doc in docs:
                tipo = doc['tipo']
                icon = {'pdf': '🔴', 'docx': '🔵', 'txt': '🟡'}.get(tipo, '⚪')
                size_str = f"{doc['size']/1024:.1f} KB"
                resumido = doc['nombre'] in st.session_state.doc_resumenes
                tiene_error = 'error' in doc
                n_chars = doc['chars']

                if tiene_error:
                    status_html = f'<span class="pill-err">❌ Error extracción</span>'
                elif n_chars == 0:
                    status_html = '<span class="pill-warn">⚠️ Sin texto</span>'
                elif resumido:
                    status_html = '<span class="pill-ok">✅ Resumido</span>'
                else:
                    status_html = '<span style="color:#8b949e;font-size:.7rem;">○ Sin resumen</span>'

                st.markdown(f"""
                <div class="doc-card {tipo}">
                  <div class="doc-title">{icon} {doc['nombre']}</div>
                  <div class="doc-meta">{size_str} · {n_chars:,} chars &nbsp; {status_html}</div>
                  {'<div style="font-size:.7rem;color:#ef4444;margin-top:4px;">⚠ ' + doc.get("error","")[:120] + '</div>' if tiene_error else ''}
                </div>
                """, unsafe_allow_html=True)

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if api_ok and st.button("🤖 Resumir todos", use_container_width=True):
                    docs_sin_resumen = [d for d in docs if d['nombre'] not in st.session_state.doc_resumenes and d['chars'] > 0]
                    if not docs_sin_resumen:
                        st.info("Todos los documentos ya están resumidos.")
                    else:
                        for doc in docs_sin_resumen:
                            nombre = doc['nombre']
                            texto = st.session_state.doc_textos.get(nombre, "")
                            with st.spinner(f"Resumiendo {nombre}..."):
                                resumen, err = generar_resumen(nombre, texto)
                            if err:
                                st.error(f"❌ {nombre}: {err}")
                            else:
                                st.session_state.doc_resumenes[nombre] = resumen
                        st.rerun()
            with col_btn2:
                if st.button("🗑️ Limpiar todo", use_container_width=True):
                    st.session_state.documentos_cargados = []
                    st.session_state.doc_textos = {}
                    st.session_state.doc_resumenes = {}
                    st.session_state.chat_history = []
                    st.rerun()
        else:
            st.markdown("""
            <div class="card" style="text-align:center;padding:28px;border-style:dashed;">
              <div style="font-size:2rem;">📄</div>
              <div style="font-size:.8rem;color:#8b949e;margin-top:6px;">Sin documentos cargados</div>
            </div>
            """, unsafe_allow_html=True)

    # ─── Panel derecho: resúmenes + chat ───
    with col_der:
        tab_res, tab_chat, tab_texto = st.tabs(["📋 RESÚMENES", "💬 CHAT CON DOCS", "📝 TEXTO EXTRAÍDO"])

        # Tab: Resúmenes
        with tab_res:
            if not docs:
                st.markdown('<div style="color:#8b949e;text-align:center;padding:40px;">Carga documentos para ver sus resúmenes</div>', unsafe_allow_html=True)
            else:
                for doc in docs:
                    nombre = doc['nombre']
                    tipo   = doc['tipo']
                    icon   = {'pdf': '🔴', 'docx': '🔵', 'txt': '🟡'}.get(tipo, '⚪')
                    n_chars = doc['chars']

                    with st.expander(f"{icon} {nombre}", expanded=True):
                        resumen = st.session_state.doc_resumenes.get(nombre)

                        if resumen:
                            # Renderizar el markdown de la IA nativamente
                            st.markdown(resumen)
                        elif n_chars == 0:
                            st.markdown("""
                            <div class="error-box">
                              ⚠️ Este documento no tiene texto extraíble (PDF escaneado o imagen).<br>
                              No es posible generar un resumen sin texto.
                            </div>
                            """, unsafe_allow_html=True)
                        else:
                            col_r1, col_r2 = st.columns([3, 1])
                            with col_r1:
                                st.markdown(f'<span style="color:#8b949e;font-size:.85rem;">Sin resumen · {n_chars:,} chars disponibles</span>', unsafe_allow_html=True)
                            with col_r2:
                                if api_ok:
                                    btn_key = f"res_{nombre.replace('.','_').replace(' ','_')}"
                                    if st.button("🤖 Resumir", key=btn_key, use_container_width=True):
                                        texto = st.session_state.doc_textos.get(nombre, "")
                                        with st.spinner(f"Generando resumen de {nombre}..."):
                                            resumen_gen, err = generar_resumen(nombre, texto)
                                        if err:
                                            st.error(f"❌ {err}")
                                        else:
                                            st.session_state.doc_resumenes[nombre] = resumen_gen
                                            st.rerun()
                                else:
                                    st.markdown('<span class="pill-warn">Requiere API key</span>', unsafe_allow_html=True)

        # Tab: Chat
        with tab_chat:
            if not docs:
                st.markdown('<div style="color:#8b949e;text-align:center;padding:40px;">Carga documentos para chatear sobre ellos</div>', unsafe_allow_html=True)
            elif not api_ok:
                st.markdown('<div class="error-box">⚙️ Configura tu API key en Configuración para usar el chat.</div>', unsafe_allow_html=True)
            else:
                historial = st.session_state.chat_history

                # ── Historial de mensajes ────────────────
                if not historial:
                    st.markdown("""
                    <div class="card" style="text-align:center;padding:28px;border-style:dashed;">
                      <div style="font-size:1.5rem;">💬</div>
                      <div style="font-size:.85rem;color:#8b949e;margin-top:8px;">
                        Haz una pregunta sobre los documentos cargados
                      </div>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    for msg in historial:
                        if msg['role'] == 'user':
                            st.markdown(f"""
                            <div class="chat-label" style="text-align:right;">Tú</div>
                            <div class="chat-msg-user">{msg['content']}</div>
                            """, unsafe_allow_html=True)
                        else:
                            st.markdown('<div class="chat-label">⚡ HACK14 IA</div>', unsafe_allow_html=True)
                            # Renderizar respuesta IA como Markdown nativo
                            st.markdown(msg['content'])

                st.markdown('<div class="h14-divider"></div>', unsafe_allow_html=True)

                # ── Sugerencias rápidas ──────────────────
                sugerencias = [
                    "¿Cuál es el tema principal?",
                    "Resume los puntos clave",
                    "¿Qué datos numéricos menciona?",
                    "¿Cuáles son las conclusiones?",
                ]
                st.markdown('<div style="font-size:.72rem;color:#8b949e;font-family:Space Mono,monospace;margin-bottom:6px;">SUGERENCIAS RÁPIDAS</div>', unsafe_allow_html=True)
                cols_sug = st.columns(len(sugerencias))
                for i, sug in enumerate(sugerencias):
                    with cols_sug[i]:
                        if st.button(sug, key=f"sug_{i}", use_container_width=True):
                            # Enviar directamente sin pasar por text_area
                            with st.spinner("Consultando documentos..."):
                                resp, err = responder_pregunta(
                                    sug,
                                    st.session_state.doc_textos,
                                    st.session_state.chat_history
                                )
                            if err:
                                st.error(f"❌ {err}")
                            else:
                                st.session_state.chat_history.append({'role': 'user',      'content': sug})
                                st.session_state.chat_history.append({'role': 'assistant', 'content': resp})
                                st.rerun()

                # ── Input de pregunta ────────────────────
                # IMPORTANTE: NO usar 'value' con 'key' en text_area (conflicto Streamlit)
                pregunta = st.text_area(
                    "Tu pregunta",
                    placeholder="Escribe tu pregunta sobre los documentos...",
                    height=90,
                    label_visibility="collapsed",
                    key="chat_input_area"
                )

                col_send, col_clear = st.columns([3, 1])
                with col_send:
                    enviar = st.button("➤ Enviar pregunta", use_container_width=True, key="btn_enviar_chat")
                with col_clear:
                    if st.button("🗑️ Limpiar chat", use_container_width=True, key="btn_limpiar_chat"):
                        st.session_state.chat_history = []
                        st.rerun()

                if enviar:
                    p = pregunta.strip()
                    if not p:
                        st.warning("Escribe una pregunta primero.")
                    else:
                        docs_con_texto = {n: t for n, t in st.session_state.doc_textos.items() if t}
                        if not docs_con_texto:
                            st.error("❌ Ningún documento tiene texto extraíble. Verifica la pestaña 'Texto Extraído'.")
                        else:
                            with st.spinner("Consultando documentos con IA..."):
                                resp, err = responder_pregunta(
                                    p,
                                    docs_con_texto,
                                    st.session_state.chat_history
                                )
                            if err:
                                st.error(f"❌ {err}")
                            else:
                                st.session_state.chat_history.append({'role': 'user',      'content': p})
                                st.session_state.chat_history.append({'role': 'assistant', 'content': resp})
                                st.rerun()

        # Tab: Texto extraído
        with tab_texto:
            if not docs:
                st.markdown('<div style="color:#8b949e;text-align:center;padding:40px;">Carga documentos para ver el texto extraído</div>', unsafe_allow_html=True)
            else:
                doc_sel = st.selectbox(
                    "Selecciona documento",
                    [d['nombre'] for d in docs],
                    label_visibility="visible"
                )
                if doc_sel:
                    texto = st.session_state.doc_textos.get(doc_sel, "")
                    st.markdown(f'<div style="font-size:.75rem;font-family:Space Mono,monospace;color:#8b949e;margin-bottom:8px;">{len(texto):,} caracteres extraídos</div>', unsafe_allow_html=True)
                    st.text_area("Texto extraído", value=texto[:10000], height=450, label_visibility="collapsed", disabled=True)
                    if len(texto) > 10000:
                        st.markdown(f'<div class="pill-warn">Mostrando primeros 10,000 de {len(texto):,} caracteres</div>', unsafe_allow_html=True)

                    txt_bytes = texto.encode('utf-8')
                    st.download_button(
                        "💾 Descargar texto completo",
                        data=txt_bytes,
                        file_name=f"{doc_sel}_texto.txt",
                        mime="text/plain"
                    )
